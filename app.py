import streamlit as st
import ee
import folium
from streamlit_folium import folium_static
import json
import pandas as pd
import requests
from datetime import datetime
from fpdf import FPDF
from docx import Document
import os
from supabase import create_client, Client

# --- 1. CONFIGURACIÓN E INICIALIZACIÓN ---
st.set_page_config(page_title="BioCore Intelligence V5", layout="wide")

@st.cache_resource
def init_db():
    return create_client(st.secrets["connections"]["supabase"]["url"], st.secrets["connections"]["supabase"]["key"])

supabase = init_db()

def iniciar_gee():
    if not ee.data.is_initialized():
        creds = json.loads(st.secrets["gee"]["json"])
        ee_creds = ee.ServiceAccountCredentials(creds['client_email'], key_data=creds['private_key'])
        ee.Initialize(ee_creds)

# --- 2. MOTOR DE DOCUMENTOS (PDF PROFESIONAL + WORD EDITABLE) ---

def generar_documentos(r):
    proyecto = r.get('proyecto', 'Proyecto_BioCore')
    # Protección total contra valores Nulos (None)
    v_ndsi = float(r.get('ndwi_ndsi') or 0)
    v_radar = float(r.get('radar_vv') or 0)
    v_temp = float(r.get('temp_suelo') or 0)
    v_savi = float(r.get('savi') or 0)
    
    # Lógica de Diagnóstico Técnico
    color = (200, 0, 0) if v_ndsi < 0.35 else (0, 100, 0)
    est = "ALERTA: PERDIDA DE COBERTURA" if v_ndsi < 0.35 else "CONTROL: ESTABILIDAD"
    diag_txt = (f"El indice detectado ({v_ndsi:.3f}) indica una degradacion de la masa criosferica. "
                f"Riesgo tecnico-legal detectado." if v_ndsi < 0.35 else 
                f"El indice ({v_ndsi:.3f}) confirma estabilidad. Cumplimiento de parametros RCA.")

    # A. GENERAR PDF (Banner Azul Marino)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_fill_color(20, 50, 80)
    pdf.rect(0, 0, 210, 40, 'F')
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("helvetica", "B", 16)
    pdf.cell(0, 20, f"AUDITORIA AMBIENTAL - {str(proyecto).upper()}", align="C", ln=1)
    pdf.set_font("helvetica", "I", 10)
    pdf.cell(0, 5, "Responsable Tecnica: Loreto Campos | BioCore Intelligence", align="C", ln=1)
    
    pdf.ln(25); pdf.set_text_color(0, 0, 0); pdf.set_font("helvetica", "B", 12)
    pdf.cell(0, 10, "DIAGNOSTICO TECNICO DE ALTA MONTAÑA", ln=1)
    
    pdf.set_fill_color(*color); pdf.set_text_color(255, 255, 255); pdf.set_font("helvetica", "B", 10)
    pdf.cell(0, 8, f" ESTATUS: {est}", ln=1, fill=True)
    
    pdf.ln(5); pdf.set_text_color(0, 0, 0); pdf.set_font("helvetica", "", 10)
    pdf.multi_cell(0, 7, f"{diag_txt}\n\nAnalisis Satelital:\n- NDSI: {v_ndsi:.3f}\n- Radar VV: {v_radar:.2f} dB\n- SAVI: {v_savi:.3f}", border=1)
    
    pdf_path = f"Reporte_{proyecto}.pdf"
    pdf.output(pdf_path)

    # B. GENERAR WORD (Editable)
    doc = Document()
    doc.add_heading(f"INFORME DE AUDITORIA: {proyecto}", 0)
    doc.add_paragraph("BioCore Intelligence - Reporte Tecnico Editable")
    doc.add_heading("DIAGNOSTICO", level=1)
    p = doc.add_paragraph(); p.add_run(f"ESTATUS ACTUAL: {est}").bold = True
    doc.add_paragraph(diag_txt)
    doc_path = f"Reporte_{proyecto}_Editable.docx"
    doc.save(doc_path)
    
    return pdf_path, doc_path

def enviar_telegram_pack(pdf_p, doc_p, proyecto):
    token = st.secrets['telegram']['token']
    chat_id = st.secrets['telegram']['chat_id']
    for f_path in [pdf_p, doc_p]:
        with open(f_path, "rb") as f:
            requests.post(f"https://api.telegram.org/bot{token}/sendDocument",
                         data={"chat_id": chat_id, "caption": f"🛡️ BioCore Intelligence: {proyecto}"},
                         files={"document": f})

# --- 3. ANALISIS ESPACIAL (SENTINEL-1 & SENTINEL-2) ---

def dibujar_mapa(coords):
    js = json.loads(coords)
    pts = js['coordinates'][0] if 'coordinates' in js else js
    loc = [[float(p[1]), float(p[0])] for p in pts]
    m = folium.Map(location=loc[0], zoom_start=15, tiles='https://mt1.google.com/vt/lyrs=y&x={x}&y={y}&z={z}', attr='Google Hybrid')
    folium.Polygon(loc, color="yellow", weight=2, fill=True, fill_opacity=0.1).add_to(m)
    return m

def ejecutar_auditoria_dual(p):
    iniciar_gee()
    geom = ee.Geometry.Polygon(json.loads(p['Coordenadas'])['coordinates'])
    
    # Optico S2
    img_s2 = ee.ImageCollection('COPERNICUS/S2_SR_HARMONIZED').filterBounds(geom).sort('CLOUDY_PIXEL_PERCENTAGE').first()
    ndsi = img_s2.normalizedDifference(['B3','B11']).reduceRegion(ee.Reducer.mean(), geom, 30).getInfo().get('nd', 0)
    savi = img_s2.expression('((B8-B4)/(B8+B4+0.5))*1.5', {'B8':img_s2.select('B8'),'B4':img_s2.select('B4')}).reduceRegion(ee.Reducer.mean(), geom, 30).getInfo().get('constant', 0)
    
    # Radar S1 (Atraviesa nubes)
    img_s1 = ee.ImageCollection('COPERNICUS/S1_GRD').filterBounds(geom).filter(ee.Filter.listContains('transmitterReceiverPolarisation', 'VV')).first()
    radar = img_s1.select('VV').reduceRegion(ee.Reducer.mean(), geom, 30).getInfo().get('VV', 0)
    
    # Temperatura MODIS
    t_img = ee.ImageCollection("MODIS/061/MOD11A1").filterBounds(geom).sort('system:time_start', False).first()
    temp = t_img.select('LST_Day_1km').multiply(0.02).subtract(273.15).reduceRegion(ee.Reducer.mean(), geom, 1000).getInfo().get('LST_Day_1km', 0)

    supabase.table("historial_reportes").insert({
        "proyecto": p['Proyecto'], "ndwi_ndsi": ndsi, "savi": savi, "radar_vv": radar, 
        "temp_suelo": temp, "validado_por_admin": False, "estado": "PENDIENTE"
    }).execute()

# --- 4. INTERFAZ STREAMLIT ---

tab1, tab2 = st.tabs(["🚀 Vigilancia Activa", "📊 Centro de Revision"])

with tab1:
    proyectos = supabase.table("usuarios").select("*").execute().data
    for p in proyectos:
        st.markdown(f"### 📍 Proyecto: {p['Proyecto']}")
        col_map, col_act = st.columns([3, 1])
        with col_map: folium_static(dibujar_mapa(p['Coordenadas']), width=950, height=450)
        with col_act:
            if st.button(f"🚀 Iniciar Auditoria", key=f"btn_{p['Proyecto']}"):
                with st.spinner("Analizando Optico + Radar..."):
                    ejecutar_auditoria_dual(p)
                    st.success("Borrador generado en Pestaña 2")

with tab2:
    st.subheader("📋 Control de Calidad y Envios Oficiales")
    
    if st.button("🗑️ Limpiar todos los borradores"):
        supabase.table("historial_reportes").delete().eq("validado_por_admin", False).execute()
        st.rerun()

    pendientes = supabase.table("historial_reportes").select("*").eq("validado_por_admin", False).execute().data
    if pendientes:
        for r in pendientes:
            # Proteccion contra Nulos para la visualizacion
            v_ndsi = float(r.get('ndwi_ndsi') or 0)
            v_radar = float(r.get('radar_vv') or 0)
            
            with st.expander(f"🔍 Revisar: {r['proyecto']} ({r['created_at'][:16]})"):
                st.write(f"**NDSI:** {v_ndsi:.3f} | **Radar:** {v_radar:.2f} dB")
                
                c1, c2, c3 = st.columns(3)
                with c1:
                    if st.button("🚀 Enviar Pack Oficial", key=f"env_{r['id']}", type="primary"):
                        with st.spinner("Generando PDF y Word..."):
                            pdf, doc = generar_documentos(r)
                            enviar_telegram_pack(pdf, doc, r['proyecto'])
                            supabase.table("historial_reportes").update({"validado_por_admin": True, "estado": "ENVIADO"}).eq("id", r['id']).execute()
                            os.remove(pdf); os.remove(doc)
                            st.success("Pack enviado."); st.rerun()
                
                with c2:
                    if st.button("📩 Reenvio Manual", key=f"man_{r['id']}"):
                        pdf, doc = generar_documentos(r)
                        enviar_telegram_pack(pdf, doc, r['proyecto'])
                        os.remove(pdf); os.remove(doc)
                        st.info("Reenviado a Telegram (Sin validar).")

                with c3:
                    if st.button("🗑️ Eliminar", key=f"del_{r['id']}"):
                        supabase.table("historial_reportes").delete().eq("id", r['id']).execute()
                        st.rerun()

    st.divider()
    st.markdown("#### 📥 Historial de Auditorias Finalizadas")
    finalizados = supabase.table("historial_reportes").select("*").eq("validado_por_admin", True).execute().data
    if finalizados:
        st.dataframe(pd.DataFrame(finalizados)[['created_at', 'proyecto', 'ndwi_ndsi', 'estado']])
